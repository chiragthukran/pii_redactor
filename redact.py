"""
Main script. Reads the input docx paragraph by paragraph, runs every
detector on each paragraph's text, resolves overlaps between detectors,
replaces each match with a consistent fake value, and writes a new docx.

Run: python3 redact.py input.docx output.docx
"""

import sys
from docx import Document

from detectors.regex_detectors import (
    find_emails, find_phones, find_ips, find_ssns, find_credit_cards, find_dobs,
)
from detectors.ner_detectors import find_people, find_companies, find_addresses
from faker_map import get_fake


# order matters here - this is priority when two matches overlap.
# addresses first because they are the widest match and contain stuff
# that the person/company detectors will also try to flag (see the
# "Village Birdewadi" tagged as PERSON issue found during testing).
# structured regex ones (email/phone/etc) are pretty much always
# correct so they take priority over the fuzzier NER ones too.
DETECTOR_ORDER = [
    ("email", find_emails),
    ("phone", find_phones),
    ("ip", find_ips),
    ("ssn", find_ssns),
    ("credit_card", find_credit_cards),
    ("dob", find_dobs),
    ("address", find_addresses),
    ("person", find_people),
    ("company", find_companies),
]


def resolve_overlaps(all_matches):
    """all_matches is a list of (category, start, end, text), already in
    priority order. keeps a match unless it overlaps with a
    higher-priority one that was already kept."""
    kept = []
    for category, start, end, text in all_matches:
        overlap = False
        for _, ks, ke, _ in kept:
            if start < ke and end > ks:  # ranges overlap
                overlap = True
                break
        if not overlap:
            kept.append((category, start, end, text))
    return kept


def redact_paragraph_text(text):
    all_matches = []
    for category, fn in DETECTOR_ORDER:
        for start, end, matched_text in fn(text):
            all_matches.append((category, start, end, matched_text))

    kept = resolve_overlaps(all_matches)
    # replace right to left so earlier offsets don't shift
    kept.sort(key=lambda m: m[1], reverse=True)

    redacted = text
    count_by_type = {}
    for category, start, end, matched_text in kept:
        fake_value = get_fake(category, matched_text)
        redacted = redacted[:start] + fake_value + redacted[end:]
        count_by_type[category] = count_by_type.get(category, 0) + 1

    return redacted, count_by_type


def redact_docx(input_path, output_path):
    doc = Document(input_path)
    total_counts = {}

    def process_paragraph(paragraph):
        if not paragraph.text.strip():
            return
        new_text, counts = redact_paragraph_text(paragraph.text)
        for k, v in counts.items():
            total_counts[k] = total_counts.get(k, 0) + v
        if new_text != paragraph.text:
            # docx splits text across multiple runs (formatting, spell
            # check markers etc) so a straight text replace can lose
            # formatting. simplest reliable fix: put all the new text in
            # the first run and clear the rest. loses run-level
            # formatting differences within a paragraph but keeps the
            # doc readable and valid, which is the main goal here.
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for r in paragraph.runs[1:]:
                    r.text = ""
            else:
                paragraph.text = new_text

    for p in doc.paragraphs:
        process_paragraph(p)

    # prospectus has a lot of content inside tables too (financial
    # statements, contact tables etc) so need to walk those separately
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)

    doc.save(output_path)
    return total_counts


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("usage: python3 redact.py input.docx output.docx")
        sys.exit(1)

    counts = redact_docx(sys.argv[1], sys.argv[2])
    print("done. redaction counts by type:")
    for k, v in sorted(counts.items()):
        print(f"  {k}: {v}")
