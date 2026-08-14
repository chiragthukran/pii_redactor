"""
Flask app - single entrypoint, Vercel's Python runtime auto-detects this
because it's at api/index.py and defines a module level `app`.

Same detection/redaction logic as the CLI version (redact.py in the main
project), just wrapped to work on an uploaded file in memory instead of
reading/writing to disk, since a serverless function doesn't have
persistent disk storage between requests.
"""

import os
import sys
import io

sys.path.insert(0, os.path.dirname(__file__))

from flask import Flask, request, send_file, jsonify
from docx import Document

from detectors.regex_detectors import (
    find_emails, find_phones, find_ips, find_ssns, find_credit_cards, find_dobs,
)
from detectors.ner_detectors import find_people, find_companies, find_addresses
from faker_map import get_fake

app = Flask(__name__)

DETECTOR_ORDER = [
    ("email", find_emails),
    ("phone", find_phones),
    ("ip", find_ips),
    ("ssn", find_ssns),
    ("credit_card", find_credit_cards),
    ("dob", find_dobs),
    ("address", find_addresses),
    ("person", find_people),
    ("company", find_companies),
]


def resolve_overlaps(all_matches):
    kept = []
    for category, start, end, text in all_matches:
        overlap = False
        for _, ks, ke, _ in kept:
            if start < ke and end > ks:
                overlap = True
                break
        if not overlap:
            kept.append((category, start, end, text))
    return kept


def redact_paragraph_text(text):
    all_matches = []
    for category, fn in DETECTOR_ORDER:
        for start, end, matched_text in fn(text):
            all_matches.append((category, start, end, matched_text))

    kept = resolve_overlaps(all_matches)
    kept.sort(key=lambda m: m[1], reverse=True)

    redacted = text
    count_by_type = {}
    for category, start, end, matched_text in kept:
        fake_value = get_fake(category, matched_text)
        redacted = redacted[:start] + fake_value + redacted[end:]
        count_by_type[category] = count_by_type.get(category, 0) + 1

    return redacted, count_by_type


def redact_docx_stream(file_stream):
    """Same as redact_docx() in the CLI version but works on an in-memory
    file object instead of a path, since python-docx's Document() and
    doc.save() both accept file-like objects, not just paths."""
    doc = Document(file_stream)
    total_counts = {}

    def process_paragraph(paragraph):
        if not paragraph.text.strip():
            return
        new_text, counts = redact_paragraph_text(paragraph.text)
        for k, v in counts.items():
            total_counts[k] = total_counts.get(k, 0) + v
        if new_text != paragraph.text:
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for r in paragraph.runs[1:]:
                    r.text = ""
            else:
                paragraph.text = new_text

    for p in doc.paragraphs:
        process_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)

    out_stream = io.BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream, total_counts


@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})


@app.route("/api/redact", methods=["POST"])
def redact_endpoint():
    if "file" not in request.files:
        return jsonify({"error": "no file uploaded - form field must be named 'file'"}), 400

    f = request.files["file"]
    if not f.filename.lower().endswith(".docx"):
        return jsonify({"error": "only .docx files are supported"}), 400

    try:
        input_stream = io.BytesIO(f.read())
        output_stream, counts = redact_docx_stream(input_stream)
    except Exception as e:
        # deliberately returning the exception message - this is a small
        # internal tool, not a public api, being able to see what broke
        # matters more here than hiding error details
        return jsonify({"error": f"failed to process file: {e}"}), 500

    download_name = "redacted_" + f.filename
    response = send_file(
        output_stream,
        as_attachment=True,
        download_name=download_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    # counts get sent back as a header rather than a separate json call -
    # keeps this to one request instead of two. exposed explicitly below
    # so the frontend fetch() call is actually allowed to read it.
    response.headers["X-Redaction-Counts"] = ",".join(f"{k}:{v}" for k, v in counts.items())
    response.headers["Access-Control-Expose-Headers"] = "X-Redaction-Counts"
    return response


# local dev only - Vercel doesn't call this, it imports `app` directly
if __name__ == "__main__":
    app.run(debug=True, port=5000)
